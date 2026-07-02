#!/usr/bin/env python3
"""
Brief Outline Generator — generate-brief.py

Renders a content OUTLINE (not a finished brief) as a .docx file.
The output is a skeleton — section headings and short topic prompts — that a
writer fills in with their own conclusions, numbers, and prose.

Usage:
    python generate-brief.py --config brief-config.json

Config JSON shape:
{
  "title": "How Do Platform Teams Implement Cloud Disaster Recovery",
  "focus_keyword": "cloud disaster recovery",
  "focus_keyword_volume": "2,400",
  "domain_url": "https://firefly.ai",
  "word_count_range": "1500-2000",
  "target_intent": "Informational",
  "target_product": "Firefly",               # optional
  "archetype": "how_to",                     # optional: listicle | comparison | how_to | concept
  "meta_title": "Cloud Disaster Recovery: A Platform Team Guide",  # optional; fetched from URL
  "meta_description": "Learn how platform teams implement cloud disaster recovery ...",  # optional
  "secondary_keywords": [
    {"keyword": "disaster recovery plan", "volume": "1,900"}
  ],
  "output_path": "/mnt/user-data/outputs/outline-xxx.docx",
  "outline": [
    {
      "heading": "H2",
      "title": "Introduction",
      "rules": ["short topic prompt 1", "short topic prompt 2"],
      "subsections": [
        {"heading": "H3", "title": "...", "rules": [...], "subsections": []}
      ]
    },
    {
      "heading": "H2",
      "title": "FAQs",
      "rules": [
        "How do I install a Claude skill?",
        "How is a Claude skill different from a regular prompt?"
      ]
    }
  ]
}

Schema notes:
- `topic_summary`, `directives`, `visual`, and `faqs` fields are no longer
  rendered. If present in the config they are silently ignored — kept for
  backward compatibility with old configs, but they produce nothing.
- `domain_context` is no longer rendered in the metadata table. The skill
  uses it upstream to inform generation, but it does not appear in the doc.
- The FAQs section uses the same `rules` shape as every other section: a list
  of question strings as topic prompts. The writer drafts the actual answers.
- The URL Slug in the metadata table is derived from `focus_keyword` (e.g.
  `cloud disaster recovery` → `cloud-disaster-recovery`). The output filename
  is derived from `title` instead — these are intentionally separate.
- `meta_title` and `meta_description` are optional string fields. When present
  they appear as "Meta Title (50-60 chars)" and "Meta Description (150-160 chars)"
  rows in the metadata table, directly above the Keywords section. If absent,
  the rows render as empty strings.
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Colour helpers ───────────────────────────────────────────────────────────

def rgb(hex_str):
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ─── XML helpers — schema-order-safe ─────────────────────────────────────────

def _set_tbl_width(table, width_dxa):
    """Set table width by updating the existing w:tblW element (always present)."""
    tblPr = table._tbl.find(qn('w:tblPr'))
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.insert(0, tblW)
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')


def _set_cell_width(cell, width_dxa):
    """Update the existing w:tcW element inside tcPr (always present after table creation)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.insert(0, tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')


def _set_cell_borders(cell, color='CCCCCC', size=4):
    """Insert w:tcBorders after w:tcW (schema order: tcW → tcBorders → shd).
    Uses w:start/w:end (OOXML 2012 strict) instead of w:left/w:right."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'start', 'bottom', 'end'):  # start/end not left/right
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color.lstrip('#'))
        tcBorders.append(el)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is not None:
        tcPr.insert(list(tcPr).index(tcW) + 1, tcBorders)
    else:
        tcPr.append(tcBorders)


def _set_cell_bg(cell, hex_color):
    """Insert w:shd after w:tcBorders (schema order: tcBorders → shd)."""
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    # Insert after tcBorders if present, else append
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        idx = list(tcPr).index(tcBorders)
        tcPr.insert(idx + 1, shd)
    else:
        tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    # OOXML strict order: top, start, bottom, end (not left/right)
    for side, val in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


# Required tcPr child order per OOXML strict schema
_TCPR_ORDER = [
    'cnfStyle', 'tcW', 'gridSpan', 'hMerge', 'vMerge',
    'tcBorders', 'shd', 'noWrap', 'tcMar', 'textDirection',
    'tcFitText', 'vAlign', 'hideMark', 'headers',
]

def _reorder_tcpr(cell):
    """Re-sort tcPr children into schema-required order. Call after cell.merge()."""
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        return
    children = list(tcPr)
    def _rank(el):
        tag = el.tag.split('}')[1]
        try:
            return _TCPR_ORDER.index(tag)
        except ValueError:
            return len(_TCPR_ORDER)
    children.sort(key=_rank)
    for child in list(tcPr):
        tcPr.remove(child)
    for child in children:
        tcPr.append(child)


def _set_cell_valign(cell, align='top'):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _style_cell(cell, width_dxa, bg_hex, border_color='CCCCCC',
                margins=(80, 80, 120, 120)):
    _set_cell_width(cell, width_dxa)
    _set_cell_borders(cell, border_color)
    _set_cell_bg(cell, bg_hex)
    _set_cell_margins(cell, *margins)
    _set_cell_valign(cell, 'top')


def _set_col_span(cell, span):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:gridSpan'))
    if existing is not None:
        tcPr.remove(existing)
    gridSpan = OxmlElement('w:gridSpan')
    gridSpan.set(qn('w:val'), str(span))
    # gridSpan must come AFTER tcW in schema order (cnfStyle? tcW? gridSpan?)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is not None:
        tcPr.insert(list(tcPr).index(tcW) + 1, gridSpan)
    else:
        tcPr.insert(0, gridSpan)


def _set_para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        pPr.remove(spacing)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)


def _apply_list_numbering(para, num_id, level=0):
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn('w:numPr'))
    if existing is not None:
        pPr.remove(existing)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(numId_el)
    # numPr must be before pStyle/jc etc — insert at front of pPr
    pPr.insert(0, numPr)


def _fix_zoom(doc):
    """Add required w:percent attribute to w:zoom element in settings."""
    zoom = doc.settings.element.find('.//' + qn('w:zoom'))
    if zoom is not None and not zoom.get(qn('w:percent')):
        zoom.set(qn('w:percent'), '100')


# ─── Numbering ────────────────────────────────────────────────────────────────

def _add_numbering_def(doc, char='•', indent_left=720, hanging=360):
    """
    Add an abstractNum + num pair to the numbering part.
    Returns the numId (int) to pass to _apply_list_numbering.
    """
    num_part = doc.part.numbering_part
    num_el = num_part._element

    # Count existing abstractNums and nums
    abstract_nums = num_el.findall(qn('w:abstractNum'))
    abstract_id = len(abstract_nums)

    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(abstract_id))

    ml_type = OxmlElement('w:multiLevelType')
    ml_type.set(qn('w:val'), 'hybridMultilevel')
    abstract_num.append(ml_type)

    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')

    for tag, val in [('w:start', '1'), ('w:numFmt', 'bullet')]:
        el = OxmlElement(tag)
        el.set(qn('w:val'), val)
        lvl.append(el)

    lvl_text = OxmlElement('w:lvlText')
    lvl_text.set(qn('w:val'), char)
    lvl.append(lvl_text)

    lvl_jc = OxmlElement('w:lvlJc')
    lvl_jc.set(qn('w:val'), 'left')
    lvl.append(lvl_jc)

    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(indent_left))
    ind.set(qn('w:hanging'), str(hanging))
    pPr.append(ind)
    lvl.append(pPr)

    abstract_num.append(lvl)

    # abstractNum must come before num elements
    existing_nums = num_el.findall(qn('w:num'))
    if existing_nums:
        num_el.insert(list(num_el).index(existing_nums[0]), abstract_num)
    else:
        num_el.append(abstract_num)

    # num element
    num_id = len(existing_nums) + 1
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(abstract_id))
    num.append(abstract_ref)
    num_el.append(num)

    return num_id


# ─── Validation ───────────────────────────────────────────────────────────────

def validate(cfg):
    errors, warnings = [], []
    if not cfg.get('title', '').strip():
        errors.append('title: must not be empty')
    elif len(cfg['title']) > 70:
        warnings.append(f"title: {len(cfg['title'])} chars — Google may truncate above 60")
    if not cfg.get('focus_keyword', '').strip():
        errors.append('focus_keyword: must not be empty')
    if not cfg.get('domain_url', '').strip():
        errors.append('domain_url: must not be empty')
    elif not re.match(r'^https?://', cfg['domain_url'], re.I):
        errors.append('domain_url: must start with http:// or https://')
    if not cfg.get('word_count_range', '').strip():
        errors.append('word_count_range: must not be empty')
    elif not re.match(r'^\d+-\d+$', cfg['word_count_range'].strip()):
        errors.append('word_count_range: e.g. 1500-2000')
    valid_intents = {'informational', 'commercial', 'transactional', 'navigational'}
    if not cfg.get('target_intent', '').strip():
        errors.append('target_intent: must not be empty')
    elif cfg['target_intent'].lower() not in valid_intents:
        errors.append('target_intent: Informational / Commercial / Transactional / Navigational')
    return errors, warnings


# ─── Slug ─────────────────────────────────────────────────────────────────────

def to_slug(title):
    s = re.sub(r'[^a-z0-9\s-]', '', title.lower()).strip()
    return re.sub(r'-+', '-', re.sub(r'\s+', '-', s)).strip('-')


# ─── Audience ─────────────────────────────────────────────────────────────────

AUDIENCE_PATTERNS = [
    (r'devops|cicd|pipeline|terraform|kubernetes|k8s|helm|argo',
     ['DevOps Engineers', 'Platform Engineers', 'SREs']),
    (r'platform|infrastructure|cloud|disaster recovery|\bdr\b|rto|rpo|resilience|failover',
     ['Platform Engineers', 'DevOps Engineers', 'Cloud Architects']),
    (r'frontend|react|vue|angular|css|\bui\b|\bux\b|nextjs|svelte',
     ['Frontend Developers', 'Full-Stack Engineers']),
    (r'backend|api|database|sql|microservices|rest|graphql|grpc',
     ['Backend Developers', 'Full-Stack Engineers', 'Software Architects']),
    (r'security|iam|auth|zero trust|compliance|soc2|devsecops',
     ['Security Engineers', 'DevSecOps']),
    (r'\bai\b|\bml\b|llm|model|embedding|vector|langchain',
     ['ML Engineers', 'AI Developers', 'Data Scientists']),
    (r'open.?source|contributor|library|sdk|cli|npm|pypi',
     ['Open-Source Contributors', 'Software Developers']),
]

def infer_audience(title, keyword):
    text = f'{title} {keyword}'.lower()
    seen, roles = set(), []
    for pattern, role_list in AUDIENCE_PATTERNS:
        if re.search(pattern, text):
            for r in role_list:
                if r not in seen:
                    seen.add(r); roles.append(r)
    return ', '.join(roles) if roles else 'Software Developers, Full-Stack Engineers, DevOps Engineers'


# ─── Paragraph builders ───────────────────────────────────────────────────────

COLORS = {
    'blue_hdr':  'D5E8F0',
    'blue_cell': 'EAF4FB',
    'blue_bg':   'EFF6FF',
    'amber_hdr': 'FEF3C7',
    'amber_row': 'FFFBEB',
    'amber_vol': 'FFF9E6',
    'white':     'FFFFFF',
}

TW = 9360   # total table width DXA
CK = 2000   # key column
CV = 5360   # value column
CS = 2000   # search volume column


def _run(para, text, bold=False, italic=False, color='1F2937', size_pt=11):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Arial'
    run.font.size = Pt(size_pt)
    run.font.color.rgb = rgb(color)
    return run


def add_h1(doc, text):
    para = doc.add_paragraph(style='Heading 1')
    _set_para_spacing(para, before=360, after=240)
    _run(para, text, bold=True, color='1F2937', size_pt=20)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph(style='Heading 2')
    _set_para_spacing(para, before=320, after=160)
    _run(para, '[H2]  ', bold=True, color='B0BEC5', size_pt=10)
    _run(para, text,    bold=True, color='1D4ED8', size_pt=15)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph(style='Heading 3')
    _set_para_spacing(para, before=220, after=120)
    _run(para, '[H3]  ', bold=True, color='D1D5DB', size_pt=9)
    _run(para, text,    bold=True, color='374151', size_pt=13)
    return para


def add_spacer(doc):
    para = doc.add_paragraph()
    _set_para_spacing(para, before=40, after=40)
    return para


def add_bullet(doc, text, num_id):
    para = doc.add_paragraph()
    _set_para_spacing(para, before=0, after=80)
    _apply_list_numbering(para, num_id)
    _run(para, text, color='1F2937', size_pt=11)
    return para


# add_visual_note and add_faq_item removed — `visual` and `faqs` fields are no longer rendered.


# ─── Tables ───────────────────────────────────────────────────────────────────

def build_metadata_table(doc, meta_rows, focus_kw, focus_vol, secondary_kws):
    # Rows: top header (Field|Value) + metadata rows + banner + KW sub-header + focus + secondaries
    n_rows = 1 + len(meta_rows) + 1 + 1 + 1 + len(secondary_kws)
    table = doc.add_table(rows=n_rows, cols=3)
    _set_tbl_width(table, TW)

    row_idx = 0

    # ── Top header row: Field | Value (Value spans cols 2+3, no volume label here)
    hk = table.cell(row_idx, 0)
    _style_cell(hk, CK, COLORS['blue_hdr'])
    _run(hk.paragraphs[0], 'Field', bold=True, color='1F2937', size_pt=11)

    hv1 = table.cell(row_idx, 1)
    hv2 = table.cell(row_idx, 2)
    hv = hv1.merge(hv2)
    _style_cell(hv, CV + CS, COLORS['blue_hdr'])
    _reorder_tcpr(hv)
    _run(hv.paragraphs[0], 'Value', bold=True, color='1F2937', size_pt=11)
    row_idx += 1

    # ── Standard metadata rows (value spans cols 2+3)
    for item in meta_rows:
        c0 = table.cell(row_idx, 0)
        _style_cell(c0, CK, COLORS['blue_cell'])
        _run(c0.paragraphs[0], item['key'], bold=True, color='1F2937', size_pt=11)

        c1 = table.cell(row_idx, 1)
        c2 = table.cell(row_idx, 2)
        merged = c1.merge(c2)
        _style_cell(merged, CV + CS, COLORS['white'])
        _reorder_tcpr(merged)
        _run(merged.paragraphs[0], item['value'] or '', color='1F2937', size_pt=11)
        row_idx += 1

    # ── Keyword section divider (banner)
    d0 = table.cell(row_idx, 0)
    d1 = table.cell(row_idx, 1)
    d2 = table.cell(row_idx, 2)
    div = d0.merge(d1).merge(d2)
    _style_cell(div, TW, COLORS['blue_hdr'])
    _reorder_tcpr(div)
    _run(div.paragraphs[0], 'Keywords & Search Volumes (Tool: Ahrefs)',
         bold=True, color='1F2937', size_pt=11)
    row_idx += 1

    # ── Keyword sub-header: Type | Keyword | USA Search Volume ↓
    for i, (label, width) in enumerate(zip(
        ['Type', 'Keyword', 'USA Search Volume \u2193'],
        [CK, CV, CS]
    )):
        c = table.cell(row_idx, i)
        _style_cell(c, width, COLORS['blue_cell'])
        p = c.paragraphs[0]
        if i == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, label, bold=True, color='1F2937', size_pt=10)
    row_idx += 1

    # ── Focus keyword row
    fk0 = table.cell(row_idx, 0)
    fk1 = table.cell(row_idx, 1)
    fk2 = table.cell(row_idx, 2)
    _style_cell(fk0, CK, COLORS['blue_cell'])
    _run(fk0.paragraphs[0], 'Focus Keyword', bold=True, color='1F2937', size_pt=11)
    _style_cell(fk1, CV, COLORS['white'])
    _run(fk1.paragraphs[0], focus_kw, color='1F2937', size_pt=11)
    _style_cell(fk2, CS, COLORS['amber_vol'])
    vp = fk2.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(vp, focus_vol or 'N/A', color='92400E', size_pt=11)
    row_idx += 1

    # ── Secondary keyword rows
    for i, kw in enumerate(secondary_kws):
        kw_text  = kw if isinstance(kw, str) else kw.get('keyword', '')
        vol_text = 'N/A' if isinstance(kw, str) else kw.get('volume', 'N/A')
        s0 = table.cell(row_idx, 0)
        s1 = table.cell(row_idx, 1)
        s2 = table.cell(row_idx, 2)
        _style_cell(s0, CK, COLORS['blue_cell'])
        _run(s0.paragraphs[0], 'Secondary KWs' if i == 0 else '',
             bold=(i == 0), color='1F2937', size_pt=11)
        _style_cell(s1, CV, COLORS['white'])
        _run(s1.paragraphs[0], kw_text, color='1F2937', size_pt=11)
        _style_cell(s2, CS, COLORS['amber_vol'])
        svp = s2.paragraphs[0]
        svp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(svp, vol_text, color='92400E', size_pt=11)
        row_idx += 1

    return table


# add_topic_summary and add_directives_box removed — `topic_summary` and
# `directives` fields are no longer rendered. Outlines don't carry abstracts
# or instruction boxes; the bullets themselves carry the direction.


# ─── Section rendering ────────────────────────────────────────────────────────

def _heading_for(heading_str, force_h3=False):
    """Map heading string to renderer. H4 falls through to H3 styling with a label."""
    if force_h3:
        return add_h3
    h = (heading_str or '').upper()
    if h == 'H4':
        # H4 is rare (listicle items). Use H3 styling — the [H4] label prefix
        # is added inside the heading text itself by the caller if needed.
        return add_h3
    if h == 'H3':
        return add_h3
    return add_h2


def render_section(doc, section, bullet_num_id, force_h3=False):
    """Render a section. The `dir_num_id` parameter is gone — directives are no longer rendered."""
    render_heading = _heading_for(section.get('heading'), force_h3=force_h3)
    render_heading(doc, section['title'])

    # Outline bullets
    for r in section.get('rules', []):
        add_bullet(doc, r, bullet_num_id)
    if section.get('rules'):
        add_spacer(doc)

    # Subsections
    for sub in section.get('subsections', []):
        render_section(doc, sub, bullet_num_id, force_h3=True)

    # NOTE: `topic_summary`, `directives`, `visual`, and `faqs` fields are
    # intentionally ignored. Outlines don't carry pre-written abstracts,
    # instruction boxes, visual placeholders, or pre-written FAQ answers.
    # If a config still includes them, they produce nothing.


# ─── Main ─────────────────────────────────────────────────────────────────────

def generate_brief(cfg):
    errors, warnings = validate(cfg)
    if errors:
        print('❌ Validation failed:')
        for e in errors:
            print(f'  - {e}')
        sys.exit(1)
    for w in warnings:
        print(f'⚠️  {w}')

    url_slug  = to_slug(cfg['focus_keyword'])   # shown in the metadata table
    file_slug = to_slug(cfg['title'])           # used for the output filename
    audience  = infer_audience(cfg['title'], cfg['focus_keyword'])
    focus_vol = cfg.get('focus_keyword_volume') or 'N/A'
    archetype = cfg.get('archetype', '').strip() or 'unspecified'

    # NOTE: domain_context is intentionally not rendered. It informs the upstream
    # outline-generation step but does not appear in the rendered metadata table.

    secondary_kws = [
        k if isinstance(k, dict) else {'keyword': k, 'volume': 'N/A'}
        for k in cfg.get('secondary_keywords', [])
    ]

    meta_rows = [
        {'key': 'Title',                            'value': cfg['title']},
        {'key': 'URL Slug',                         'value': url_slug},
        {'key': 'Word Count',                       'value': cfg['word_count_range']},
        {'key': 'Target Intent',                    'value': cfg['target_intent']},
        {'key': 'Target Audience',                  'value': audience},
        {'key': 'Meta Title (50-60 chars)',         'value': cfg.get('meta_title', '')},
        {'key': 'Meta Description (150-160 chars)', 'value': cfg.get('meta_description', '')},
    ]

    # ── Build document
    doc = Document()

    # Fix python-docx zoom element bug
    _fix_zoom(doc)

    # Page: US Letter, 1" margins
    sec = doc.sections[0]
    sec.page_width  = Twips(12240)
    sec.page_height = Twips(15840)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin  = sec.bottom_margin = Inches(1)

    # Default + heading styles
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    for name, size, color in [
        ('Heading 1', 20, '1F2937'),
        ('Heading 2', 15, '1D4ED8'),
        ('Heading 3', 13, '374151'),
    ]:
        s = doc.styles[name]
        s.font.name = 'Arial'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = rgb(color)

    # List numbering — only bullets. The directives arrow-list is gone.
    bullet_num_id = _add_numbering_def(doc, char='•', indent_left=720, hanging=360)

    # 1. Metadata + keyword volume table
    build_metadata_table(doc, meta_rows, cfg['focus_keyword'], focus_vol, secondary_kws)
    add_spacer(doc)
    add_spacer(doc)

    # 2. H1 title
    add_h1(doc, cfg['title'])
    add_spacer(doc)

    # 3. Outline sections
    for section in cfg.get('outline', []):
        render_section(doc, section, bullet_num_id, force_h3=False)

    # Save — note: filename uses 'outline-' prefix (this is an outline generator).
    # Ensure the parent directory exists so the script works outside the sandboxed
    # Claude environment where /mnt/user-data/outputs/ is pre-created.
    output_path = cfg.get('output_path') or f'/mnt/user-data/outputs/outline-{file_slug}.docx'
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    kw_summary = ', '.join(
        f"{k.get('keyword', '?')} ({k.get('volume', 'N/A')})" for k in secondary_kws
    )
    print(f'✅ Outline generated: {output_path}')
    print(f'   Archetype:     {archetype}')
    print(f'   Slug (URL):    {url_slug}')
    print(f'   Slug (file):   {file_slug}')
    print(f'   Audience:      {audience}')
    print(f'   Focus KW:      {cfg["focus_keyword"]} (vol: {focus_vol})')
    print(f'   Secondary KWs: {kw_summary or "none"}')
    print(f'   Sections:      {len(cfg.get("outline", []))} top-level')
    return {'success': True, 'path': output_path}


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--config', required=True, help='Path to brief-config.json')
    args = parser.parse_args()
    with open(args.config) as f:
        cfg = json.load(f)
    generate_brief(cfg)